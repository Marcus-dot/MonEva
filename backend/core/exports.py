"""
Word document export generator for MonEva project reports.
Uses python-docx to produce professionally formatted .docx files.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from django.utils import timezone


# Brand colours
GREEN_DARK = RGBColor(0x1A, 0x4D, 0x2E)
GREEN_LIGHT = RGBColor(0xE8, 0xF5, 0xE9)
GREY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ── Helpers ──────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via raw XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _heading_row(table, *labels):
    """Style the first row of a table as a dark green header."""
    row = table.rows[0]
    for i, label in enumerate(labels):
        cell = row.cells[i]
        cell.text = label
        _set_cell_bg(cell, '1A4D2E')
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)


def _add_section_heading(doc, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREEN_DARK
    # Bottom border line
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A4D2E')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_kv(doc, label: str, value: str):
    """Add a key: value line."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    run_label = para.add_run(f'{label}: ')
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_value = para.add_run(str(value) if value else '—')
    run_value.font.size = Pt(10)
    run_value.font.color.rgb = GREY


def _status_badge(status: str) -> str:
    """Map status codes to readable labels."""
    mapping = {
        'PENDING': 'Pending',
        'IN_PROGRESS': 'In Progress',
        'COMPLETED': 'Completed',
        'PLANNING': 'Planning',
        'ACTIVE': 'Active',
        'ON_HOLD': 'On Hold',
        'DRAFT': 'Draft',
        'SUBMITTED': 'Submitted',
        'APPROVED': 'Approved',
        'PAID': 'Paid',
        'REJECTED': 'Rejected',
        'OPEN': 'Open',
        'INVESTIGATING': 'Investigating',
        'RESOLVED': 'Resolved',
        'CLOSED': 'Closed',
        'EXPIRED': 'Expired',
        'SUSPENDED': 'Suspended',
    }
    return mapping.get(status, status)


# ── Main Report Builder ───────────────────────────────────────────────────────

def generate_project_report_docx(project) -> bytes:
    """
    Generate a full project status report as a .docx file.

    Args:
        project: Project model instance (with prefetched contracts, milestones)

    Returns:
        bytes: The .docx file content
    """
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Cover / Title block ──────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('PROJECT STATUS REPORT')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = GREEN_DARK

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(project.name)
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = GREY

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(
        f'Generated: {timezone.now().strftime("%d %B %Y")}  |  KGLCDC MonEva System'
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = GREY

    doc.add_paragraph()  # spacer

    # ── 1. Project Overview ──────────────────────────────────────────────────
    _add_section_heading(doc, '1. Project Overview')
    _add_kv(doc, 'Project Name', project.name)
    _add_kv(doc, 'Type', project.get_type_display())
    _add_kv(doc, 'Status', _status_badge(project.status))
    _add_kv(doc, 'Start Date', str(project.start_date))
    _add_kv(doc, 'End Date', str(project.end_date))
    if project.description:
        _add_kv(doc, 'Description', project.description)
    if project.chiefdom if hasattr(project, 'chiefdom') else None:
        _add_kv(doc, 'Location', project.chiefdom)

    # SDGs
    sdgs = project.sdgs.all()
    if sdgs.exists():
        _add_kv(doc, 'SDGs', ', '.join(f'{s.code}: {s.name}' for s in sdgs))

    # ── 2. Contracts & Milestones ────────────────────────────────────────────
    contracts = project.contracts.prefetch_related('milestones', 'claims').all()

    _add_section_heading(doc, '2. Contracts & Milestones')

    if not contracts.exists():
        doc.add_paragraph('No contracts recorded for this project.').italic = True
    else:
        for contract in contracts:
            # Contract header row
            contract_label = contract.contract_number or f'Contract {str(contract.id)[:8]}'
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(10)
            run = para.add_run(f'{contract_label}  —  {contract.contractor.name}')
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = GREEN_DARK

            _add_kv(doc, 'Type', contract.get_contract_type_display())
            _add_kv(doc, 'Status', _status_badge(contract.status))
            _add_kv(doc, 'Value', f'{contract.currency} {contract.total_value:,.2f}')
            _add_kv(doc, 'Period', f'{contract.start_date} to {contract.end_date}')
            if contract.chiefdom:
                _add_kv(doc, 'Chiefdom', contract.chiefdom)

            milestones = contract.milestones.all()
            if milestones.exists():
                tbl = doc.add_table(rows=1, cols=4)
                tbl.style = 'Table Grid'
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                _heading_row(tbl, 'Milestone', 'Due Date', 'Target %', 'Status')

                for ms in milestones:
                    row_cells = tbl.add_row().cells
                    row_cells[0].text = ms.title
                    row_cells[1].text = str(ms.due_date)
                    row_cells[2].text = f'{ms.target_percent}%'
                    row_cells[3].text = _status_badge(ms.status)
                    for cell in row_cells:
                        cell.paragraphs[0].runs[0].font.size = Pt(9)

                doc.add_paragraph()  # spacer after table
            else:
                p = doc.add_paragraph('No milestones recorded.')
                p.runs[0].italic = True
                p.runs[0].font.size = Pt(9)

    # ── 3. Financial Summary ─────────────────────────────────────────────────
    _add_section_heading(doc, '3. Financial Summary')

    from finance.models import PaymentClaim
    claims = PaymentClaim.objects.filter(contract__project=project)

    if not claims.exists():
        doc.add_paragraph('No payment claims recorded.').runs[0].italic = True
    else:
        total_claimed = sum(c.amount for c in claims)
        approved = claims.filter(status='APPROVED')
        paid = claims.filter(status='PAID')
        total_approved = sum(c.amount for c in approved)
        total_paid = sum(c.amount for c in paid)

        _add_kv(doc, 'Total Claimed', f'ZMW {total_claimed:,.2f}')
        _add_kv(doc, 'Total Approved', f'ZMW {total_approved:,.2f}')
        _add_kv(doc, 'Total Paid', f'ZMW {total_paid:,.2f}')
        _add_kv(doc, 'Pending Claims', str(claims.filter(status='SUBMITTED').count()))

        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'
        _heading_row(tbl, 'Claim Date', 'Amount (ZMW)', 'Prepared By', 'Status')

        for claim in claims.order_by('claim_date'):
            row_cells = tbl.add_row().cells
            row_cells[0].text = str(claim.claim_date)
            row_cells[1].text = f'{claim.amount:,.2f}'
            row_cells[2].text = claim.prepared_by.get_full_name() if claim.prepared_by else '—'
            row_cells[3].text = _status_badge(claim.status)
            for cell in row_cells:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()

    # ── 4. Grievances Summary ────────────────────────────────────────────────
    _add_section_heading(doc, '4. Grievances Summary')

    from grievances.models import Grievance
    grievances = Grievance.objects.filter(project=project)

    if not grievances.exists():
        doc.add_paragraph('No grievances recorded for this project.').runs[0].italic = True
    else:
        _add_kv(doc, 'Total Grievances', str(grievances.count()))
        _add_kv(doc, 'Open', str(grievances.filter(status='OPEN').count()))
        _add_kv(doc, 'Investigating', str(grievances.filter(status='INVESTIGATING').count()))
        _add_kv(doc, 'Resolved', str(grievances.filter(status='RESOLVED').count()))
        _add_kv(doc, 'Closed', str(grievances.filter(status='CLOSED').count()))

    # ── 5. Indicators Summary ────────────────────────────────────────────────
    _add_section_heading(doc, '5. Indicators (M&E)')

    try:
        from indicators.models import LogFrameNode
        indicators = LogFrameNode.objects.filter(project=project, node_type='INDICATOR')

        if not indicators.exists():
            doc.add_paragraph('No indicators configured for this project.').runs[0].italic = True
        else:
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = 'Table Grid'
            _heading_row(tbl, 'Indicator', 'Baseline', 'Target')

            for ind in indicators:
                row_cells = tbl.add_row().cells
                row_cells[0].text = ind.title or ind.description or '—'
                row_cells[1].text = str(ind.baseline_value) if hasattr(ind, 'baseline_value') and ind.baseline_value is not None else '—'
                row_cells[2].text = str(ind.target_value) if hasattr(ind, 'target_value') and ind.target_value is not None else '—'
                for cell in row_cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)

            doc.add_paragraph()
    except Exception:
        doc.add_paragraph('Indicators data unavailable.').runs[0].italic = True

    # ── Footer note ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph(
        f'This report was automatically generated by the MonEva M&E System on '
        f'{timezone.now().strftime("%d %B %Y at %H:%M")}. '
        'For queries contact your system administrator.'
    )
    note.runs[0].font.size = Pt(8)
    note.runs[0].font.color.rgb = GREY
    note.paragraph_format.space_before = Pt(20)

    # ── Serialise to bytes ────────────────────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
